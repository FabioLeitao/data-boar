"""
ATS/SLI Export Helper — converte .md para DOCX e TXT puro.
Uso: chamado pelo export-ats-sli.ps1 via uv run python.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ---------------------------------------------------------------------------
# Markdown → plain TXT (strip all formatting)
# ---------------------------------------------------------------------------
_MD_STRIP_PATTERNS = [
    (re.compile(r"^#{1,6}\s+"), ""),  # headings
    (re.compile(r"\*\*(.+?)\*\*"), r"\1"),  # bold
    (re.compile(r"\*(.+?)\*"), r"\1"),  # italic
    (re.compile(r"`(.+?)`"), r"\1"),  # inline code
    (re.compile(r"```[\w]*\n(.*?)```", re.DOTALL), r"\1"),  # code blocks
    (re.compile(r"^\s*[-*+]\s+", re.M), "• "),  # bullet lists
    (re.compile(r"^\s*\d+\.\s+", re.M), ""),  # numbered lists
    (re.compile(r"!\[.*?\]\(.*?\)"), ""),  # images
    (re.compile(r"\[(.+?)\]\(.*?\)"), r"\1"),  # links → text
    (re.compile(r"^\|.*\|$", re.M), lambda m: re.sub(r"\|", "  ", m.group())),  # tables
    (re.compile(r"^[-|:]+$", re.M), ""),  # table separators
    (re.compile(r"^>\s+", re.M), "  "),  # blockquotes
    (re.compile(r"---+"), ""),  # horizontal rules
    (re.compile(r"\n{3,}"), "\n\n"),  # excess blank lines
]


def md_to_txt(text: str) -> str:
    for pattern, repl in _MD_STRIP_PATTERNS:
        if callable(repl):
            text = pattern.sub(repl, text)
        else:
            text = pattern.sub(repl, text)
    return text.strip()


# ---------------------------------------------------------------------------
# Markdown → DOCX (basic, paragraph-level)
# ---------------------------------------------------------------------------
def md_to_docx(text: str, out_path: Path, title: str = "") -> None:
    if not HAS_DOCX:
        raise ImportError("python-docx not available")

    doc = Document()

    # Title style
    if title:
        t = doc.add_heading(title, level=0)
        t.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    code_block = False
    code_lines: list[str] = []

    def flush_code(lines: list[str]) -> None:
        if lines:
            p = doc.add_paragraph("\n".join(lines))
            p.style = "Normal"
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            if not code_block:
                code_block = True
                code_lines = []
            else:
                flush_code(code_lines)
                code_block = False
                code_lines = []
            i += 1
            continue

        if code_block:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=min(level, 4))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            p = doc.add_paragraph()
            p.add_run("─" * 60)
            i += 1
            continue

        # Table row (simplified — merge all cells into one paragraph)
        if re.match(r"^\|", line):
            cells = [c.strip() for c in line.strip("|").split("|")]
            # separator row
            if all(re.match(r"^[-:]+$", c.replace(" ", "")) for c in cells if c):
                i += 1
                continue
            p = doc.add_paragraph("  |  ".join(c for c in cells if c))
            p.style = "Normal"
            for run in p.runs:
                run.font.size = Pt(9.5)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Pt(20)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # Bullet list
        m = re.match(r"^\s*[-*+]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(md_inline(m.group(1)), style="List Bullet")
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(md_inline(m.group(1)), style="List Number")
            i += 1
            continue

        # Checkbox (- [ ] / - [x])
        m = re.match(r"^\s*-\s+\[([ x])\]\s+(.*)", line)
        if m:
            done = m.group(1) == "x"
            doc.add_paragraph(
                f"{'☑' if done else '☐'} {m.group(2)}", style="List Bullet"
            )
            i += 1
            continue

        # Empty line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph(md_inline(line))
        i += 1

    if code_block and code_lines:
        flush_code(code_lines)

    doc.save(str(out_path))


def md_inline(text: str) -> str:
    """Strip inline markdown formatting for docx runs."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.*?\)", r"\1", text)
    return text


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main() -> None:
    parser = argparse.ArgumentParser(description="Convert .md to DOCX and/or TXT")
    parser.add_argument("input", type=Path, help="Source .md file")
    parser.add_argument(
        "--out-dir",
        type=Path,
        required=True,
        help="Output directory root (must contain docx/ and txt/ subdirs)",
    )
    parser.add_argument(
        "--formats", nargs="+", choices=["docx", "txt"], default=["docx", "txt"]
    )
    args = parser.parse_args()

    src: Path = args.input
    if not src.exists():
        print(f"ERROR: {src} not found", file=sys.stderr)
        sys.exit(1)

    text = src.read_text(encoding="utf-8")
    stem = src.stem  # filename without extension

    if "txt" in args.formats:
        out = args.out_dir / "txt" / f"{stem}.txt"
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(md_to_txt(text), encoding="utf-8")
        print(f"TXT → {out}")

    if "docx" in args.formats:
        out = args.out_dir / "docx" / f"{stem}.docx"
        out.parent.mkdir(parents=True, exist_ok=True)
        title = re.sub(r"^#\s+", "", text.split("\n")[0]).strip()
        md_to_docx(text, out, title=title)
        print(f"DOCX → {out}")


if __name__ == "__main__":
    main()
