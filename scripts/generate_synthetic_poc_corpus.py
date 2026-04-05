#!/usr/bin/env python3
"""
generate_synthetic_poc_corpus.py
================================
Generates a synthetic test corpus for Data Boar POC validation.

Covers seven test scenarios:
  1. happy          -- clear PII in plain-text formats (should be found)
  2. unhappy        -- PII with OCR noise, encoding quirks (should be found, harder)
  3. catastrophic   -- nested archives, password-protected zips (may be missed)
  4. false_positive -- data that LOOKS like PII but is invalid (should NOT trigger)
  5. manual_review  -- ambiguous / partial data (flag for human review)
  6. stego          -- CPF/RG hidden in image LSB (NOT found without stego module)
  7. extensions     -- one file per supported extension, all containing a CPF

Usage:
  uv run python scripts/generate_synthetic_poc_corpus.py
  uv run python scripts/generate_synthetic_poc_corpus.py --scenario happy,stego
  uv run python scripts/generate_synthetic_poc_corpus.py --output /tmp/poc_corpus

Collaborator note:
  After generating, point Data Boar at each sub-folder and compare findings
  against the expected results in EXPECTED.txt (each sub-folder) and
  docs/TESTING_POC_GUIDE.md (full validation checklist).
"""

from __future__ import annotations

import argparse
import csv
import io
import json
import sqlite3
import tarfile
import textwrap
import zipfile
from pathlib import Path
from typing import Callable

# ---------------------------------------------------------------------------
# Synthetic PII (deterministic, never real persons)
# ---------------------------------------------------------------------------
_CPFS = [
    "123.456.789-09",
    "987.654.321-00",
    "111.222.333-96",
    "000.000.001-91",
    "529.982.247-25",
]
_CNPJS = ["11.222.333/0001-81", "00.000.000/0001-91", "12.345.678/0001-95"]
_RGS = ["12.345.678-9", "98.765.432-1", "00.111.222-3"]
_NAMES = [
    "Ana Paula Souza",
    "Carlos Eduardo Lima",
    "Fernanda Beatriz Costa",
    "Joao Roberto Almeida",
    "Maria Oliveira Santos",
]
_EMAILS = [
    "ana.souza@example-test.com",
    "carlos.lima@demo.invalid",
    "f.costa@poc-databoar.test",
]
_PHONES = ["(11) 99999-0001", "+55 21 98888-0002", "0800 123 4567"]
_DATES = ["15/03/1985", "1990-07-22", "01/01/1970"]
_ADDRS = [
    "Rua das Flores, 123, Sao Paulo - SP, CEP 01234-567",
    "Av. Brasil, 4500, Rio de Janeiro - RJ",
]

EXPECTED: dict[str, str] = {
    "1_happy": "DEVE ENCONTRAR -- PII em claro, sem ofuscacao",
    "2_unhappy": "DEVE ENCONTRAR -- mas pode requerer OCR ou tolerancia a ruido",
    "3_catastrophic": "PODE NAO ENCONTRAR -- dados em arquivos aninhados ou com senha",
    "4_false_positive": "NAO DEVE ENCONTRAR -- strings similares a PII mas invalidas",
    "5_manual_review": "DEVE SINALIZAR PARA REVISAO MANUAL -- dados parcialmente mascarados",
    "6_stego": "NAO DEVE ENCONTRAR sem modulo estego -- CPF em LSB de imagem PNG",
    "7_extensions": "DEVE ENCONTRAR em todos os formatos suportados",
}


def _p(lst: list[str], i: int = 0) -> str:
    return lst[i % len(lst)]


def _w(path: Path, content: str | bytes, enc: str = "utf-8") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if isinstance(content, bytes):
        path.write_bytes(content)
    else:
        path.write_text(content, encoding=enc)


# ---------------------------------------------------------------------------
# Scenario 1 - Happy path
# ---------------------------------------------------------------------------
def gen_scenario_1(base: Path) -> None:
    out = base / "1_happy"

    _w(
        out / "employees.txt",
        textwrap.dedent(f"""\
        RELATORIO DE FUNCIONARIOS -- FICTICIO -- APENAS PARA TESTES POC
        Nome: {_p(_NAMES, 0)}  CPF: {_p(_CPFS, 0)}  RG: {_p(_RGS, 0)}
        Email: {_p(_EMAILS, 0)}  Tel: {_p(_PHONES, 0)}
        Nasc: {_p(_DATES, 0)}   End: {_p(_ADDRS, 0)}
        Nome: {_p(_NAMES, 1)}  CPF: {_p(_CPFS, 1)}  CNPJ: {_p(_CNPJS, 0)}
    """),
    )

    buf = io.StringIO()
    csv.writer(buf).writerows(
        [["nome", "cpf", "rg", "email"]]
        + [[_p(_NAMES, i), _p(_CPFS, i), _p(_RGS, i), _p(_EMAILS, i)] for i in range(5)]
    )
    _w(out / "employees.csv", buf.getvalue())

    _w(
        out / "employees.json",
        json.dumps(
            [
                {"nome": _p(_NAMES, i), "cpf": _p(_CPFS, i), "rg": _p(_RGS, i)}
                for i in range(4)
            ],
            ensure_ascii=False,
            indent=2,
        ),
    )

    try:
        from reportlab.pdfgen import canvas as rc

        c = rc.Canvas(str(out / "employees.pdf"))
        y = 780
        c.drawString(50, y, "DADOS FICTICIOS -- POC Data Boar")
        y -= 20
        for i in range(3):
            for lbl, v in [
                ("Nome", _p(_NAMES, i)),
                ("CPF", _p(_CPFS, i)),
                ("Email", _p(_EMAILS, i)),
            ]:
                c.drawString(50, y, f"{lbl}: {v}")
                y -= 15
        c.save()
    except ImportError:
        _w(out / "employees_pdf_fallback.txt", f"PDF nao gerado. CPF: {_p(_CPFS, 0)}")

    try:
        import docx as _d

        doc = _d.Document()
        doc.add_heading("Dados Ficticios POC", 0)
        for i in range(3):
            doc.add_paragraph(
                f"Nome: {_p(_NAMES, i)}\nCPF: {_p(_CPFS, i)}\nRG: {_p(_RGS, i)}\n"
            )
        doc.save(str(out / "employees.docx"))
    except ImportError:
        _w(out / "employees_docx_fallback.txt", f"DOCX nao gerado. CPF: {_p(_CPFS, 1)}")

    try:
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Funcionarios"
        ws.append(["Nome", "CPF", "RG", "Email", "Tel"])
        for i in range(5):
            ws.append(
                [
                    _p(_NAMES, i),
                    _p(_CPFS, i),
                    _p(_RGS, i),
                    _p(_EMAILS, i),
                    _p(_PHONES, i),
                ]
            )
        wb.save(str(out / "employees.xlsx"))
    except ImportError:
        _w(out / "employees_xlsx_fallback.txt", f"XLSX nao gerado. CPF: {_p(_CPFS, 2)}")

    conn = sqlite3.connect(str(out / "employees.db"))
    conn.execute(
        "CREATE TABLE IF NOT EXISTS emp (id INTEGER PRIMARY KEY,nome TEXT,cpf TEXT,rg TEXT,email TEXT)"
    )
    for i in range(5):
        conn.execute(
            "INSERT INTO emp(nome,cpf,rg,email) VALUES(?,?,?,?)",
            (_p(_NAMES, i), _p(_CPFS, i), _p(_RGS, i), _p(_EMAILS, i)),
        )
    conn.commit()
    conn.close()

    try:
        from PIL import Image, ImageDraw

        img = Image.new("RGB", (400, 120), (255, 255, 255))
        draw = ImageDraw.Draw(img)
        draw.text((10, 20), f"CPF: {_p(_CPFS, 0)}", (0, 0, 0))
        draw.text((10, 50), f"Nome: {_p(_NAMES, 0)}", (0, 0, 0))
        img.save(str(out / "id_card_visible.png"))
    except ImportError:
        _w(out / "image_fallback.txt", f"PNG nao gerado. CPF: {_p(_CPFS, 0)}")

    _w(out / "EXPECTED.txt", EXPECTED["1_happy"])
    print(f"  v  Scenario 1 (happy) -> {out}")


# ---------------------------------------------------------------------------
# Scenario 2 - Unhappy path
# ---------------------------------------------------------------------------
def gen_scenario_2(base: Path) -> None:
    out = base / "2_unhappy"
    import base64 as _b64

    _w(
        out / "ocr_noisy.txt",
        f"N0me: {_p(_NAMES, 2).replace('a', '@').replace('e', '3')}\n"
        f"CPF: {_p(_CPFS, 2).replace('.', ',')}  (possivel ruido OCR)\n"
        f"RG: {_p(_RGS, 2).replace('-', '_')}\n"
        f"Email: {_p(_EMAILS, 2).replace('@', '[at]')}\n",
    )

    _w(
        out / "latin1_encoded.txt",
        f"Nome: {_p(_NAMES, 0)}\nCPF: {_p(_CPFS, 0)}\nObservacao: dado em latin-1\n",
        enc="latin-1",
    )

    _w(
        out / "bom_utf8.csv",
        f"\ufeffNome;CPF;RG\n{_p(_NAMES, 1)};{_p(_CPFS, 1)};{_p(_RGS, 1)}\n",
        enc="utf-8-sig",
    )

    _w(
        out / "crlf_endings.txt",
        f"CPF: {_p(_CPFS, 3)}\r\nTel: {_p(_PHONES, 0)}\r\nEnd: {_p(_ADDRS, 0)}\r\n",
    )

    _w(
        out / "partial_redaction.txt",
        f"Nome: {_p(_NAMES, 0)}\n"
        f"CPF: ***.{_p(_CPFS, 0)[4:7]}.***-**  (parcialmente redactado)\n"
        f"Email: {_p(_EMAILS, 0)}\nRG: {_p(_RGS, 0)}\n",
    )

    blob = _b64.b64encode(f"CPF:{_p(_CPFS, 1)},Nome:{_p(_NAMES, 1)}".encode()).decode()
    _w(out / "base64_embedded.txt", f"campo_documento: {blob}\n# dado acima e base64\n")

    _w(out / "EXPECTED.txt", EXPECTED["2_unhappy"])
    print(f"  v  Scenario 2 (unhappy) -> {out}")


# ---------------------------------------------------------------------------
# Scenario 3 - Catastrophic
# ---------------------------------------------------------------------------
def gen_scenario_3(base: Path) -> None:
    out = base / "3_catastrophic"
    out.mkdir(parents=True, exist_ok=True)
    pii = f"DADOS FICTICIOS\nCPF: {_p(_CPFS, 4)}\nNome: {_p(_NAMES, 4)}\nRG: {_p(_RGS, 2)}\n".encode()

    # nested zip
    inner_buf = io.BytesIO()
    with zipfile.ZipFile(inner_buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("pii.txt", pii)
    with zipfile.ZipFile(out / "nested.zip", "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("inner.zip", inner_buf.getvalue())

    # password-protected zip
    with zipfile.ZipFile(out / "password_protected.zip", "w", zipfile.ZIP_STORED) as z:
        z.setpassword(b"poc-test-123")
        z.writestr("secret.txt", pii)

    # tar.gz
    with tarfile.open(str(out / "archive.tar.gz"), "w:gz") as t:
        info = tarfile.TarInfo("pii.txt")
        info.size = len(pii)
        t.addfile(info, io.BytesIO(pii))

    # tar.bz2
    with tarfile.open(str(out / "archive.tar.bz2"), "w:bz2") as t:
        info = tarfile.TarInfo("pii.txt")
        info.size = len(pii)
        t.addfile(info, io.BytesIO(pii))

    # disguised extension (text file named .jpg)
    _w(
        out / "report_2026.jpg",
        pii.decode() + "\n# Arquivo de texto mascarado como .jpg\n",
    )

    # very long line stress test
    _w(
        out / "long_line_stress.txt",
        "x" * 5000 + f" CPF: {_p(_CPFS, 0)} " + "y" * 5000 + "\n",
    )

    _w(out / "EXPECTED.txt", EXPECTED["3_catastrophic"])
    _w(
        out / "PASSWORD_HINT.txt",
        "Senha: poc-test-123\nConfiguracao: zip_password no config.yaml\n",
    )
    print(f"  v  Scenario 3 (catastrophic) -> {out}")


# ---------------------------------------------------------------------------
# Scenario 4 - False positive pressure
# ---------------------------------------------------------------------------
def gen_scenario_4(base: Path) -> None:
    import random as _r

    out = base / "4_false_positive"

    def _invalid_cpf_shaped() -> str:
        d = [_r.randint(0, 9) for _ in range(9)]
        return f"{d[0]}{d[1]}{d[2]}.{d[3]}{d[4]}{d[5]}.{d[6]}{d[7]}{d[8]}-{(d[0] + 1) % 10}{(d[1] + 1) % 10}"

    _w(
        out / "serial_numbers.txt",
        "CATALOGO -- FICÇÃO\n"
        + "\n".join(f"Serial: {_invalid_cpf_shaped()}" for _ in range(10)),
    )
    _w(out / "cnpj_shaped_refs.txt", "Ref: 00.111.222/0099-00\n" * 5)  # invalid CNPJ
    _w(
        out / "random_codes.txt",
        "\n".join(f"Cod: {_r.randint(10000000000, 99999999999)}" for _ in range(20)),
    )
    _w(
        out / "ip_addresses.txt",
        "\n".join(f"IP: 192.168.{i}.{j}" for i in range(5) for j in range(5)),
    )
    _w(
        out / "version_strings.txt",
        "\n".join(f"v: 1.{i}.{i + 1}-{i + 2}" for i in range(10)),
    )

    _w(out / "EXPECTED.txt", EXPECTED["4_false_positive"])
    print(f"  v  Scenario 4 (false_positive) -> {out}")


# ---------------------------------------------------------------------------
# Scenario 5 - Manual review triggers
# ---------------------------------------------------------------------------
def gen_scenario_5(base: Path) -> None:
    out = base / "5_manual_review"
    _w(
        out / "masked_pii.txt",
        textwrap.dedent("""\
        CPF: ***.456.789-**   (mascarado -- padrao parcial visivel)
        CPF: 123.***.***-09   (mascarado -- inicio e fim visiveis)
        RG: 12.345.***-*
        Email: a***.s***@example.com
        Tel: (11) 9****-0001
        Nome: Ana P. S.  (iniciais -- identificacao possivel com contexto)
    """),
    )
    _w(
        out / "pii_in_prose.txt",
        textwrap.dedent("""\
        O documento de CPF terminado em 09 foi verificado.
        O numero de registro e 123456789 (sem pontuacao -- validacao manual necessaria).
        O titular nasceu em quinze de marco de 1985.
    """),
    )
    _w(
        out / "foreign_pii.txt",
        textwrap.dedent("""\
        DNI: 12345678A  (Espanha -- nao e CPF brasileiro)
        SSN: 123-45-6789  (EUA -- nao e CPF)
        NIF: X1234567L  (Espanha -- estrangeiro)
    """),
    )
    _w(
        out / "anonymized_columns.csv",
        "cpf,nome,email\n[ANONIMIZADO],[ANONIMIZADO],[ANONIMIZADO]\n" * 5,
    )

    _w(out / "EXPECTED.txt", EXPECTED["5_manual_review"])
    print(f"  v  Scenario 5 (manual_review) -> {out}")


# ---------------------------------------------------------------------------
# Scenario 6 - Steganography (LSB + EXIF metadata)
# ---------------------------------------------------------------------------
def _embed_lsb(img_path: Path, secret: str) -> None:
    from PIL import Image

    img = Image.new("RGB", (200, 200), (200, 200, 200))
    pixels = list(img.getdata())
    bits = "".join(f"{ord(c):08b}" for c in secret) + "00000000"
    new_pixels = []
    for i, (r, g, b) in enumerate(pixels):
        if i < len(bits):
            r = (r & 0xFE) | int(bits[i])
        new_pixels.append((r, g, b))
    out_img = Image.new("RGB", (200, 200))
    out_img.putdata(new_pixels)
    out_img.save(str(img_path), format="PNG")


def _extract_lsb(img_path: Path) -> str:
    from PIL import Image

    pixels = list(Image.open(str(img_path)).getdata())
    bits = [str(r & 1) for r, g, b in pixels]
    chars = []
    for i in range(0, len(bits) - 8, 8):
        c = chr(int("".join(bits[i : i + 8]), 2))
        if c == "\x00":
            break
        chars.append(c)
    return "".join(chars)


def gen_scenario_6(base: Path) -> None:
    out = base / "6_stego"
    out.mkdir(parents=True, exist_ok=True)
    try:
        from PIL import Image
        from PIL.PngImagePlugin import PngInfo

        secret = f"CPF:{_p(_CPFS, 0)};Nome:{_p(_NAMES, 0)}"
        stego_path = out / "innocent_photo.png"
        _embed_lsb(stego_path, secret)
        recovered = _extract_lsb(stego_path)
        assert recovered == secret, f"LSB mismatch: {recovered!r}"

        _w(
            out / "STEGO_KEY.txt",
            f"Arquivo: innocent_photo.png\n"
            f"Dado oculto (LSB canal R): {secret}\n"
            f"Metodo: LSB -- canal R da imagem PNG (1 bit por pixel)\n"
            f"Para extrair manualmente: use stegosuite, steghide, ou a funcao _extract_lsb() neste script.\n"
            f"Verificacao OK: dado recuperado = {recovered!r}\n",
        )

        # EXIF / PNG metadata injection
        img = Image.new("RGB", (200, 200), (180, 200, 220))
        meta = PngInfo()
        meta.add_text("Comment", f"CPF:{_p(_CPFS, 1)} Nome:{_p(_NAMES, 1)}")
        meta.add_text("Author", _p(_NAMES, 1))
        img.save(str(out / "photo_with_exif_pii.png"), pnginfo=meta)

        _w(
            out / "EXPECTED.txt",
            EXPECTED["6_stego"] + "\n\n"
            "VALIDACAO MANUAL:\n"
            "1. innocent_photo.png -- CPF em LSB. Scanner padrao NAO detecta.\n"
            '   Extrair: uv run python -c "'
            "from scripts.generate_synthetic_poc_corpus import _extract_lsb;"
            "from pathlib import Path; print(_extract_lsb(Path('tests/synthetic_corpus/6_stego/innocent_photo.png')))\"\n"
            "2. photo_with_exif_pii.png -- CPF em metadado PNG (Comment). Scanner PODE detectar se ler metadata.\n",
        )

        print(f"  v  Scenario 6 (stego) -> {out} [LSB OK, recovered={recovered!r}]")

    except ImportError:
        _w(
            out / "EXPECTED.txt",
            "Pillow nao disponivel -- cenario 6 nao gerado.\nInstale: pip install pillow\n"
            + EXPECTED["6_stego"],
        )
        print(
            "  !  Scenario 6 (stego) -> Pillow indisponivel, documentado sem gerar imagem"
        )


# ---------------------------------------------------------------------------
# Scenario 7 - Extension coverage (one file per supported extension)
# ---------------------------------------------------------------------------
def gen_all_extensions(base: Path) -> None:
    out = base / "7_extensions"
    out.mkdir(parents=True, exist_ok=True)
    pii = f"CPF: {_p(_CPFS, 0)}\nNome: {_p(_NAMES, 0)}\n"
    pii_b = pii.encode()

    for ext in [
        ".txt",
        ".log",
        ".md",
        ".rst",
        ".cfg",
        ".ini",
        ".env",
        ".yml",
        ".yaml",
        ".sql",
    ]:
        _w(out / f"sample{ext}", pii)

    _w(
        out / "sample.json",
        json.dumps({"cpf": _p(_CPFS, 0), "nome": _p(_NAMES, 0)}, ensure_ascii=False),
    )
    _w(
        out / "sample.xml",
        f'<?xml version="1.0"?><r><cpf>{_p(_CPFS, 0)}</cpf><nome>{_p(_NAMES, 0)}</nome></r>',
    )
    _w(out / "sample.csv", f"cpf,nome\n{_p(_CPFS, 0)},{_p(_NAMES, 0)}\n")
    _w(out / "sample.tsv", f"cpf\tnome\n{_p(_CPFS, 0)}\t{_p(_NAMES, 0)}\n")

    with zipfile.ZipFile(out / "sample.zip", "w") as z:
        z.writestr("pii.txt", pii)
    with tarfile.open(str(out / "sample.tar.gz"), "w:gz") as t:
        i = tarfile.TarInfo("pii.txt")
        i.size = len(pii_b)
        t.addfile(i, io.BytesIO(pii_b))
    with tarfile.open(str(out / "sample.tar.bz2"), "w:bz2") as t:
        i = tarfile.TarInfo("pii.txt")
        i.size = len(pii_b)
        t.addfile(i, io.BytesIO(pii_b))

    conn = sqlite3.connect(str(out / "sample.db"))
    conn.execute("CREATE TABLE t(cpf TEXT,nome TEXT)")
    conn.execute("INSERT INTO t VALUES(?,?)", (_p(_CPFS, 0), _p(_NAMES, 0)))
    conn.commit()
    conn.close()

    try:
        import openpyxl

        wb = openpyxl.Workbook()
        wb.active.append(["cpf", "nome"])
        wb.active.append([_p(_CPFS, 0), _p(_NAMES, 0)])
        wb.save(str(out / "sample.xlsx"))
    except ImportError:
        pass
    try:
        import docx as _d

        doc = _d.Document()
        doc.add_paragraph(pii)
        doc.save(str(out / "sample.docx"))
    except ImportError:
        pass
    try:
        from reportlab.pdfgen import canvas as rc

        c = rc.Canvas(str(out / "sample.pdf"))
        c.drawString(50, 750, f"CPF: {_p(_CPFS, 0)}")
        c.save()
    except ImportError:
        pass
    try:
        from PIL import Image, ImageDraw

        img = Image.new("RGB", (300, 80), (255, 255, 255))
        draw = ImageDraw.Draw(img)
        draw.text((10, 20), f"CPF: {_p(_CPFS, 0)}", (0, 0, 0))
        img.save(str(out / "sample.png"))
        img.save(str(out / "sample.jpg"))
    except ImportError:
        pass

    _w(
        out / "EXPECTED.txt",
        "Todos os arquivos contem CPF 123.456.789-09.\n"
        "O scanner DEVE encontrar em todos os formatos suportados.\n"
        "Formatos NAO encontrados = gap de cobertura para documentar.\n",
    )
    print(f"  v  Scenario 7 (extensions) -> {out}")


# ---------------------------------------------------------------------------
# Main


# ---------------------------------------------------------------------------
# Scenario 8 - Stress / Load (OOM, large files, high concurrency corpus)
# ---------------------------------------------------------------------------
def gen_stress_load(base: Path) -> None:
    """
    Generates files designed to stress the scanner:
    - Very large text file with PII scattered at known offsets
    - Many small files (directory flood)
    - Deeply nested directory tree
    - File with millions of lines (minimal PII density)
    - Binary file with PII embedded in non-printable bytes
    All are expected to be found; OOM or timeout = reportable failure.
    """
    out = base / "8_stress_load"
    out.mkdir(parents=True, exist_ok=True)

    # Large file: 50 MB of padding with 10 CPF instances
    large = out / "large_50mb.txt"
    chunk = "x" * 1000 + "\n"
    with large.open("w", encoding="utf-8") as f:
        for i in range(50000):  # ~50 MB
            f.write(chunk)
            if i % 5000 == 0 and i > 0:
                f.write(f"CPF: {_p(_CPFS, i // 5000)}\nNome: {_p(_NAMES, i // 5000)}\n")
    print(f"    -> large file: {large} ({large.stat().st_size // 1024 // 1024} MB)")

    # Directory flood: 500 tiny files
    flood_dir = out / "directory_flood"
    flood_dir.mkdir(exist_ok=True)
    for i in range(500):
        (flood_dir / f"file_{i:04d}.txt").write_text(
            f"ref:{i}\nCPF: {_p(_CPFS, i)}\n"
            if i % 50 == 0
            else f"ref:{i}\nnada aqui\n",
            encoding="utf-8",
        )
    print(f"    -> directory flood: 500 files (10 with PII) -> {flood_dir}")

    # Deep nesting: 10 levels, PII at the bottom
    deep = out / "deep_nesting"
    current = deep
    for lvl in range(10):
        current = current / f"level_{lvl:02d}"
        current.mkdir(parents=True, exist_ok=True)
    (current / "hidden_pii.txt").write_text(
        f"CPF: {_p(_CPFS, 0)}\nNome: {_p(_NAMES, 0)}\n# 10 levels deep\n",
        encoding="utf-8",
    )
    print(f"    -> deep nesting (10 levels): {current / 'hidden_pii.txt'}")

    # High line count: 1 million lines, PII on lines 100000, 500000, 999999
    million_lines = out / "million_lines.txt"
    with million_lines.open("w", encoding="utf-8") as f:
        for i in range(1_000_000):
            if i in {100_000, 500_000, 999_999}:
                f.write(f"CPF: {_p(_CPFS, i % len(_CPFS))}\n")
            else:
                f.write(f"linha {i}\n")
    print(f"    -> million lines: {million_lines}")

    _w(
        out / "EXPECTED.txt",
        "STRESS TEST -- OBJETIVO: scanner nao deve crashar nem perder PIIs.\n"
        "Esperado: CPF encontrado em large_50mb.txt (10x), directory_flood (10 arquivos),\n"
        "deep_nesting/hidden_pii.txt, e million_lines.txt (3x).\n"
        "Falha: OOM, timeout, crash, ou PII nao encontrado.\n"
        "Metrica: tempo de scan, memoria maxima (medir com /usr/bin/time -v ou psutil).\n",
    )

    _w(
        out / "STRESS_TEST_COMMANDS.sh",
        "#!/bin/bash\n"
        "# Medir tempo e memoria do scan de stress\n"
        "/usr/bin/time -v uv run python main.py \\\n"
        "    --config config.yaml \\\n"
        "    --scan --target tests/synthetic_corpus/8_stress_load \\\n"
        "    --report 2> stress_metrics.txt\n"
        "echo 'Metricas em stress_metrics.txt'\n"
        "grep -E 'Maximum resident|Elapsed|Exit code' stress_metrics.txt\n",
    )
    print(f"  v  Scenario 8 (stress/load) -> {out}")


# ---------------------------------------------------------------------------
# Scenario 9 - Config errors (intentional misconfigs for UX/error message QA)
# ---------------------------------------------------------------------------
def gen_config_errors(base: Path) -> None:
    """
    Generates intentionally broken config files + a test script to run each.
    The goal is NOT to scan PII — it is to evaluate:
    - Quality of error messages (stdout/stderr)
    - Dashboard troubleshooting recommendations
    - Recovery / retry behavior
    Each config has a documented EXPECTED_ERROR and TROUBLESHOOT hint.
    """
    out = base / "9_config_errors"
    out.mkdir(parents=True, exist_ok=True)

    configs: list[dict] = [
        {
            "name": "wrong_db_host",
            "description": "Database host does not exist (DNS failure)",
            "config": {
                "targets": [
                    {
                        "type": "postgresql",
                        "host": "nonexistent-db.local",
                        "port": 5432,
                        "database": "testdb",
                        "user": "admin",
                        "password": "secret",
                    }
                ],
                "report": {"output_dir": "./reports"},
            },
            "expected_error": "connection refused OR DNS resolution failure",
            "troubleshoot": "Verifique se o host esta acessivel (ping / nslookup). "
            "Confirme VPN ativa se DB for interno.",
        },
        {
            "name": "wrong_db_credentials",
            "description": "Valid host but wrong username/password",
            "config": {
                "targets": [
                    {
                        "type": "postgresql",
                        "host": "localhost",
                        "port": 5432,
                        "database": "testdb",
                        "user": "wrong_user",
                        "password": "wrong_pass",
                    }
                ],
                "report": {"output_dir": "./reports"},
            },
            "expected_error": "authentication failed for user 'wrong_user'",
            "troubleshoot": "Verifique as credenciais. Use variavel de ambiente DB_PASSWORD "
            "em vez de senha em texto no config.",
        },
        {
            "name": "missing_output_dir",
            "description": "Report output_dir does not exist and cannot be created",
            "config": {
                "targets": [
                    {"type": "filesystem", "path": "./tests/synthetic_corpus/1_happy"}
                ],
                "report": {"output_dir": "/nonexistent/readonly/path"},
            },
            "expected_error": "permission denied OR directory not found",
            "troubleshoot": "Crie o diretorio manualmente ou aponte para um caminho gravavel. "
            "Em Docker: monte o volume correto.",
        },
        {
            "name": "invalid_target_type",
            "description": "Unknown connector type specified",
            "config": {
                "targets": [
                    {"type": "oracle_xyz_invalid", "host": "localhost", "port": 1521}
                ],
                "report": {"output_dir": "./reports"},
            },
            "expected_error": "unknown connector type 'oracle_xyz_invalid'",
            "troubleshoot": "Tipos validos: postgresql, mysql, mssql, oracle, mongodb, redis, "
            "filesystem. Verifique a documentacao em docs/USAGE.md.",
        },
        {
            "name": "malformed_yaml",
            "description": "Syntactically invalid YAML config",
            "raw_content": "targets:\n  - type: postgresql\n    host: localhost\n  bad yaml: [unclosed\n",
            "expected_error": "YAML parse error",
            "troubleshoot": "Valide o YAML em https://www.yamllint.com/ ou com: "
            "python -c \"import yaml; yaml.safe_load(open('config.yaml'))\"",
        },
        {
            "name": "missing_required_field",
            "description": "Config missing required 'targets' key",
            "config": {
                "report": {"output_dir": "./reports"},
            },
            "expected_error": "missing required field 'targets' in config",
            "troubleshoot": "Copie o config de exemplo: cp deploy/config.example.yaml config.yaml "
            "e edite os targets.",
        },
        {
            "name": "path_not_found",
            "description": "Filesystem target path does not exist",
            "config": {
                "targets": [
                    {"type": "filesystem", "path": "/nonexistent/data/path/12345"}
                ],
                "report": {"output_dir": "./reports"},
            },
            "expected_error": "path '/nonexistent/data/path/12345' does not exist",
            "troubleshoot": "Confirme que o caminho existe e que o usuario tem permissao de leitura. "
            "Em Docker: monte o volume com -v /seu/caminho:/data.",
        },
        {
            "name": "api_key_wrong",
            "description": "API request with wrong X-API-Key header",
            "config": {
                "targets": [
                    {"type": "filesystem", "path": "./tests/synthetic_corpus/1_happy"}
                ],
                "api": {"require_api_key": True, "api_key": "correct-key-12345"},
                "report": {"output_dir": "./reports"},
            },
            "expected_error": "HTTP 401 Unauthorized when calling API with wrong key",
            "troubleshoot": "Use X-API-Key: correct-key-12345 no header. "
            "Para testar: curl -H 'X-API-Key: wrong-key' http://localhost:8088/api/v1/scan",
            "test_curl": (
                "curl -s -o /dev/null -w '%{http_code}' "
                "-H 'X-API-Key: WRONG-KEY' http://localhost:8088/api/v1/status"
            ),
        },
    ]

    import yaml as _yaml  # may not be available; fall back to json dump

    test_script_lines = [
        "#!/bin/bash",
        "# Auto-generated: test each broken config and capture exit code + output",
        "# Usage: bash 9_config_errors/run_error_tests.sh 2>&1 | tee error_test_results.txt",
        "",
        "PASS=0; FAIL=0; SKIP=0",
        "",
    ]

    for cfg in configs:
        cfg_path = out / f"config_{cfg['name']}.yaml"
        if "raw_content" in cfg:
            _w(cfg_path, cfg["raw_content"])
        else:
            try:
                import yaml as _yaml

                _w(
                    cfg_path,
                    _yaml.dump(
                        cfg["config"], allow_unicode=True, default_flow_style=False
                    ),
                )
            except ImportError:
                _w(cfg_path, json.dumps(cfg["config"], ensure_ascii=False, indent=2))

        doc_path = out / f"doc_{cfg['name']}.txt"
        _w(
            doc_path,
            (
                f"Config: {cfg['name']}\n"
                f"Descricao: {cfg['description']}\n"
                f"Erro esperado: {cfg['expected_error']}\n"
                f"Troubleshoot: {cfg['troubleshoot']}\n"
                + (f"Teste curl: {cfg.get('test_curl', 'N/A')}\n")
            ),
        )

        name_val = cfg["name"]
        cfg_file = cfg_path.name
        scan_tgt = "./tests/synthetic_corpus/1_happy"
        test_script_lines += [
            f'echo "--- Testing: {name_val} ---"',
            f"uv run python main.py --config {cfg_file} --scan --target {scan_tgt} 2>&1 | head -20",
            "RC=$?; if [ $RC -ne 0 ]; then"
            f' echo "EXPECTED FAILURE (rc=$RC): {name_val} -- OK"; PASS=$((PASS+1));'
            f' else echo "UNEXPECTED SUCCESS: {name_val} -- REVIEW"; FAIL=$((FAIL+1)); fi',
            "",
        ]

    test_script_lines += [
        'echo ""',
        'echo "Results: PASS=$PASS  FAIL=$FAIL  SKIP=$SKIP"',
        'echo "(PASS = expected failure triggered correctly)"',
    ]

    _w(out / "run_error_tests.sh", "\n".join(test_script_lines))

    _w(
        out / "EXPECTED.txt",
        "CENARIO 9 -- CONFIG ERRORS\n"
        "Objetivo: avaliar qualidade das mensagens de erro e recomendacoes de troubleshooting.\n"
        "Cada config_*.yaml e proposital e incorreto.\n\n"
        "Para cada caso, avaliar:\n"
        "  [ ] Mensagem de erro e clara e actionable?\n"
        "  [ ] Exit code nao-zero (distingue erro de sucesso)?\n"
        "  [ ] Dashboard mostra recomendacao de troubleshooting?\n"
        "  [ ] Nenhum stacktrace interno exposto para usuario final?\n"
        "  [ ] Log tem nivel correto (ERROR vs WARNING vs INFO)?\n\n"
        "Score qualitativo (1-5 por caso):\n"
        "  5 = mensagem clara, troubleshoot acionavel, sem stacktrace, exit code correto\n"
        "  1 = crash sem mensagem, stacktrace exposto, exit 0 em erro\n",
    )
    print(f"  v  Scenario 9 (config_errors) -> {out} ({len(configs)} configs)")


# ---------------------------------------------------------------------------
_SCENARIO_MAP: dict[str, Callable[[Path], None]] = {
    "happy": gen_scenario_1,
    "unhappy": gen_scenario_2,
    "catastrophic": gen_scenario_3,
    "false_positive": gen_scenario_4,
    "manual_review": gen_scenario_5,
    "stego": gen_scenario_6,
    "extensions": gen_all_extensions,
    "stress_load": gen_stress_load,
    "config_errors": gen_config_errors,
}
ALL_SCENARIOS = list(_SCENARIO_MAP)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate synthetic POC corpus for Data Boar."
    )
    parser.add_argument(
        "--output",
        default="tests/synthetic_corpus",
        help="Output directory (default: tests/synthetic_corpus)",
    )
    parser.add_argument(
        "--scenario",
        default=",".join(ALL_SCENARIOS),
        help=f"Comma-separated scenarios. Options: {', '.join(ALL_SCENARIOS)}",
    )
    args = parser.parse_args()

    base = Path(args.output)
    base.mkdir(parents=True, exist_ok=True)
    selected = [s.strip() for s in args.scenario.split(",")]
    unknown = [s for s in selected if s not in _SCENARIO_MAP]
    if unknown:
        parser.error(f"Unknown scenarios: {unknown}")

    print("\nData Boar -- Synthetic POC Corpus Generator")
    print(f"Output:    {base.resolve()}")
    print(f"Scenarios: {selected}\n")

    for name in selected:
        _SCENARIO_MAP[name](base)

    manifest = {
        "generated_by": "generate_synthetic_poc_corpus.py",
        "scenarios": {
            name: EXPECTED.get(f"{i + 1}_{name}", "see EXPECTED.txt")
            for i, name in enumerate(ALL_SCENARIOS)
        },
        "note": "All PII is synthetic -- generated for testing only. Not real individuals.",
    }
    (base / "CORPUS_MANIFEST.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    print(f"\nManifest -> {base / 'CORPUS_MANIFEST.json'}")
    print(
        f"Next:  uv run python main.py --config config.yaml --scan --target {base.resolve()}"
    )
    print("       Compare findings against EXPECTED.txt in each sub-folder.")
    print("       See docs/TESTING_POC_GUIDE.md for the full validation checklist.\n")


if __name__ == "__main__":
    main()
